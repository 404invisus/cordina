#!/usr/bin/env python3
"""Isi salinan dokumen UAT dengan hasil pengujian nyata + screenshot."""
import json, os, shutil
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE = os.path.dirname(os.path.abspath(__file__))
SRC = '/home/ymjsty/dev/agrawork/UAT_ConnectOne_v3.docx'
OUT = '/home/ymjsty/dev/agrawork/UAT_ConnectOne_v3_HASIL-VERIFIKASI-TEKNIS.docx'

results = {r['id']: r for r in json.load(open(os.path.join(BASE, 'ALL-RESULTS.json')))}

# tabel skenario ke-N (indeks tabel di dokumen) -> daftar id skenario, berurutan
TABLE_MAP = {
    2:  ['3.1-1', '3.1-2', '3.1-3'],
    3:  ['3.2-1', '3.2-2', '3.2-3', '3.2-4', '3.2-5'],
    4:  ['3.3.1-1', '3.3.1-2'],
    5:  ['3.3.2-1', '3.3.2-2', '3.3.2-3'],
    6:  ['3.3.3-1', '3.3.3-2'],
    7:  ['3.3.4-1', '3.3.4-2'],
    8:  ['3.4.1-1', '3.4.1-2', '3.4.1-3', '3.4.1-4'],
    9:  ['3.4.2-1', '3.4.2-2'],
    10: ['3.4.3-1', '3.4.3-2', '3.4.3-3'],
    11: ['3.4.4-1', '3.4.4-2', '3.4.4-3'],
    12: ['3.4.5-1', '3.4.5-2', '3.4.5-3'],
    13: ['3.5.1-1'],
    14: ['3.5.2-1', '3.5.2-2'],
    15: ['3.5.3-1'],
    16: ['3.5.4-1'],
    17: ['3.6.1-1'],
    18: ['3.6.2-1'],
    19: ['3.6.3-1'],
    20: ['3.6.4-1', '3.6.4-2'],
    21: ['3.7.1-1', '3.7.1-2', '3.7.1-3'],
    22: ['3.7.2-1', '3.7.2-2'],
    23: ['3.8.1-1', '3.8.1-2', '3.8.1-3', '3.8.1-4'],
    24: ['3.8.2-1', '3.8.2-2'],
    25: ['3.8.3-1'],
    26: ['4.1-1', '4.1-2', '4.1-3', '4.1-4', '4.1-5'],
    27: ['4.2-1', '4.2-2', '4.2-3', '4.2-4', '4.2-5', '4.2-6'],
    28: ['4.3-1', '4.3-2'],
}

# modul -> (nama bab pada rekap, daftar id)
MODULES = [
    ('3.1 Modul Autentikasi',              ['3.1-1', '3.1-2', '3.1-3']),
    ('3.2 Modul Dashboard',                ['3.2-1', '3.2-2', '3.2-3', '3.2-4', '3.2-5']),
    ('3.3 Modul Penjadwalan Agenda',       ['3.3.1-1', '3.3.1-2', '3.3.2-1', '3.3.2-2', '3.3.2-3',
                                            '3.3.3-1', '3.3.3-2', '3.3.4-1', '3.3.4-2']),
    ('3.4 Modul Manajemen Proyek',         ['3.4.1-1', '3.4.1-2', '3.4.1-3', '3.4.1-4', '3.4.2-1', '3.4.2-2',
                                            '3.4.3-1', '3.4.3-2', '3.4.3-3', '3.4.4-1', '3.4.4-2', '3.4.4-3',
                                            '3.4.5-1', '3.4.5-2', '3.4.5-3']),
    ('3.5 Modul Manajemen Perubahan',      ['3.5.1-1', '3.5.2-1', '3.5.2-2', '3.5.3-1', '3.5.4-1']),
    ('3.6 Modul Tanda Tangan Elektronik',  ['3.6.1-1', '3.6.2-1', '3.6.3-1', '3.6.4-1', '3.6.4-2']),
    ('3.7 Sistem Notifikasi',              ['3.7.1-1', '3.7.1-2', '3.7.1-3', '3.7.2-1', '3.7.2-2']),
    ('3.8 Manajemen Pengguna & Hak Akses', ['3.8.1-1', '3.8.1-2', '3.8.1-3', '3.8.1-4',
                                            '3.8.2-1', '3.8.2-2', '3.8.3-1']),
    ('4.1 Sprint & Kanban',                ['4.1-1', '4.1-2', '4.1-3', '4.1-4', '4.1-5']),
    ('4.2 Manajemen Dokumen',              ['4.2-1', '4.2-2', '4.2-3', '4.2-4', '4.2-5', '4.2-6']),
    ('4.3 Storage / Penyimpanan File',     ['4.3-1', '4.3-2']),
]

GREEN = RGBColor(0x1B, 0x7F, 0x3B)
RED   = RGBColor(0xC0, 0x1B, 0x1B)
AMBER = RGBColor(0xB0, 0x6A, 0x00)
GREY  = RGBColor(0x55, 0x55, 0x55)


def verdict_label(v):
    if v == 'FAIL':
        return '✗ Fail', RED
    if v == 'TIDAK TERSEDIA':
        return '✗ Fail\n(fitur tidak tersedia)', RED
    if 'catatan' in v:
        return '✓ Pass\n(dengan catatan)', AMBER
    return '✓ Pass', GREEN


def set_cell(cell, text, color=None, bold=False, size=9):
    cell.text = ''
    p = cell.paragraphs[0]
    for i, line in enumerate(str(text).split('\n')):
        r = p.add_run(('\n' if i else '') + line)
        r.font.size = Pt(size)
        r.font.bold = bold
        if color:
            r.font.color.rgb = color


def add_note(cell, rec):
    """Isi baris kosong di bawah skenario dengan catatan + screenshot."""
    cell.text = ''
    p = cell.paragraphs[0]
    r = p.add_run('Hasil pengamatan: ')
    r.font.bold = True
    r.font.size = Pt(8.5)
    r2 = p.add_run(rec['notes'])
    r2.font.size = Pt(8.5)

    shots = [s for s in rec.get('shots', []) if s]
    seen = set()
    shown = 0
    for s in shots:
        path = s if os.path.isabs(s) else os.path.join(BASE, s)
        if not os.path.exists(path) or path in seen:
            continue
        seen.add(path)
        cap = cell.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        try:
            cap.add_run().add_picture(path, width=Inches(5.9))
            shown += 1
        except Exception as e:
            cap.add_run(f'[gambar tidak dapat dimuat: {os.path.basename(path)}]').font.size = Pt(8)
        lbl = cell.add_paragraph()
        lbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lr = lbl.add_run(f'Gambar bukti {shown} — {os.path.basename(path)}')
        lr.font.size = Pt(7.5)
        lr.font.italic = True
        lr.font.color.rgb = GREY
    if shown == 0:
        cell.add_paragraph().add_run('(tidak ada tangkapan layar)').font.size = Pt(8)


def main():
    shutil.copy(SRC, OUT)
    doc = Document(OUT)
    tables = doc.tables

    filled = 0
    for ti, ids in TABLE_MAP.items():
        t = tables[ti]
        for k, sid in enumerate(ids):
            rec = results.get(sid)
            if rec is None:
                continue
            scen_row = 1 + k * 2
            note_row = scen_row + 1
            if scen_row >= len(t.rows):
                continue
            label, color = verdict_label(rec['verdict'])
            set_cell(t.rows[scen_row].cells[3], label, color, bold=True, size=9)
            if note_row < len(t.rows):
                merged = t.rows[note_row].cells[0]
                for c in t.rows[note_row].cells[1:]:
                    merged = merged.merge(c)
                add_note(merged, rec)
            filled += 1
    print(f'  skenario terisi: {filled}')

    # ---- Rekapitulasi (tabel 29) ----
    rec_t = tables[29]
    tot_p = tot_f = tot_n = tot_all = 0
    for i, (name, ids) in enumerate(MODULES, start=1):
        rs = [results[x] for x in ids if x in results]
        p = sum(1 for r in rs if r['verdict'].startswith('PASS'))
        f = len(rs) - p
        tot_p += p; tot_f += f; tot_all += len(rs)
        if i < len(rec_t.rows):
            set_cell(rec_t.rows[i].cells[0], name, size=9)
            set_cell(rec_t.rows[i].cells[1], str(len(rs)), size=9)
            set_cell(rec_t.rows[i].cells[2], str(p), GREEN if p else None, size=9)
            set_cell(rec_t.rows[i].cells[3], str(f), RED if f else None, size=9)
            set_cell(rec_t.rows[i].cells[4], '0', size=9)
    last = len(MODULES) + 1
    if last < len(rec_t.rows):
        set_cell(rec_t.rows[last].cells[0], 'TOTAL', bold=True, size=9)
        set_cell(rec_t.rows[last].cells[1], str(tot_all), bold=True, size=9)
        set_cell(rec_t.rows[last].cells[2], str(tot_p), GREEN, bold=True, size=9)
        set_cell(rec_t.rows[last].cells[3], str(tot_f), RED if tot_f else None, bold=True, size=9)
        set_cell(rec_t.rows[last].cells[4], '0', bold=True, size=9)
    print(f'  rekap: {tot_all} skenario, {tot_p} pass, {tot_f} fail')

    # ---- Catatan & tindak lanjut (tabel 30) ----
    fails = [r for r in results.values() if not r['verdict'].startswith('PASS')]
    fails.sort(key=lambda r: r['id'])
    ft = tables[30]
    set_cell(ft.rows[0].cells[1],
             ('\n'.join(f'{r["id"]} — {r["scenario"]}' for r in fails) if fails else
              'Tidak ada temuan kritis yang masih terbuka.\n\n'
              'Pada putaran verifikasi pertama ditemukan 5 ketidaksesuaian. Seluruhnya telah diperbaiki '
              'dan diuji ulang hingga sesuai ekspektasi:\n'
              '  • 3.4.4-2 Filter beban kerja per anggota — filter anggota ditambahkan\n'
              '  • 3.6.2-1 Multi penanda tangan berurutan — nama parameter ke layanan TTE diperbaiki\n'
              '  • 3.8.1-4 Nonaktifkan pengguna — pemeriksaan status akun ditambahkan pada login dan rute terproteksi\n'
              '  • 4.1-3 Tambah backlog ke sprint — tombol pemicu modal ditambahkan\n'
              '  • 4.2-5 Edit dokumen — pemuatan data ke form dan metode pengiriman pembaruan diperbaiki'),
             size=9)
    set_cell(ft.rows[1].cells[1],
             '1. Cabut token lintas layanan saat akun dinonaktifkan. Saat ini penonaktifan langsung berlaku pada '
             'svc-auth, namun layanan lain memvalidasi token secara lokal sehingga token lama masih berlaku sampai '
             'kedaluwarsa (maksimal 1 jam).\n'
             '2. Selaraskan bahasa antarmuka dengan ekspektasi dokumen (antarmuka berbahasa Inggris, dokumen '
             'berbahasa Indonesia), atau sesuaikan penulisan dokumen.\n'
             '3. Beri status eksplisit pada notifikasi Telegram yang tidak terkirim karena penerima belum memiliki '
             'telegram_chat_id (saat ini tertahan berstatus "pending" tanpa keterangan).\n'
             '4. Sesuaikan langkah pengujian pada dokumen agar cocok dengan alur aplikasi (lihat berkas '
             'TEMUAN-PENGUJIAN.md bagian B3), termasuk pengisian peserta agenda, aksi Submit pada Change Request, '
             'dan pemilihan sprint pada halaman laporan.\n'
             '5. Bersihkan kode debug yang tertinggal pada AdminCalendarController dan samakan kosakata status Epic '
             'antara form, pilihan dropdown, dan nilai yang tersimpan.', size=9)
    set_cell(ft.rows[2].cells[1],
             'Butir 1 sebelum go-live; butir 2–5 sebelum pelaksanaan UAT resmi bersama BLPID', size=9)

    doc.save(OUT)
    print('  tersimpan:', OUT)


if __name__ == '__main__':
    main()
