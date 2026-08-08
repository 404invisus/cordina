"""Isi Buku Panduan ConnectOne, bagian A: sampul s.d. Bab 5 (Kalender)."""
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Inches
from build_panduan import NAVY, GREY, DIM


def cover(D):
    d = D.d
    for _ in range(3):
        d.add_paragraph()
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('[ LOGO INSTANSI ]')
    r.font.size = Pt(11)
    r.font.color.rgb = DIM
    d.add_paragraph()

    def big(text, size, bold=True, color=NAVY, after=4):
        p = d.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(after)
        r = p.add_run(text)
        r.bold = bold
        r.font.size = Pt(size)
        r.font.color.rgb = color

    big('BUKU PANDUAN PENGGUNAAN APLIKASI', 18)
    big('(USER MANUAL)', 13, bold=False, color=GREY, after=24)
    big('ConnectOne', 30, after=4)
    big('Platform Manajemen Kerja Internal BLPID', 12, bold=False, color=GREY, after=36)
    big('Balai Layanan Penghubung Identitas Digital', 12, color=NAVY, after=2)
    big('Badan Siber dan Sandi Negara', 12, color=NAVY, after=2)
    big('Tahun 2026', 11, bold=False, color=GREY)
    D.pagebreak()


def info_dokumen(D):
    D.d.add_paragraph('Informasi Dokumen', style='Heading 1')
    D.kv_table([
        ('Nama Dokumen', 'Buku Panduan Penggunaan Aplikasi ConnectOne'),
        ('Versi Dokumen', '1.0'),
        ('Versi Aplikasi', '[1.0.0]'),
        ('Tanggal Terbit', '[DD Bulan YYYY]'),
        ('Disusun oleh', '[Nama Penyusun / PT Triasta Maju Bersama]'),
        ('Diperiksa oleh', '[Nama Pemeriksa]'),
        ('Disetujui oleh', '[Nama Penyetuju]'),
        ('Klasifikasi', '[Terbatas / Internal]'),
    ])

    D.d.add_paragraph('Riwayat Revisi', style='Heading 2')
    D.table(['Versi', 'Tanggal', 'Deskripsi Perubahan', 'Penulis'], [
        ['1.0', '[tanggal]', 'Penerbitan awal', '[nama]'],
        ['[..]', '[tanggal]', '[deskripsi]', '[nama]'],
    ], widths=[0.7, 1.2, 3.1, 1.2])

    D.d.add_paragraph('Daftar Isi', style='Heading 1')
    D.p('Untuk menampilkan daftar isi otomatis di Microsoft Word: buka tab '
        '**References → Table of Contents**, lalu pilih salah satu format. '
        'Setelah dokumen diperbarui, klik kanan pada daftar isi dan pilih '
        '**Update Field → Update entire table** untuk memperbarui nomor halaman.',
        size=10, color=GREY)
    D.pagebreak()


def bab1_pendahuluan(D):
    D.h1('Pendahuluan')

    D.h2(1, 'Tujuan Dokumen')
    D.p('Dokumen ini disusun sebagai panduan bagi pengguna dalam mengoperasikan '
        'aplikasi ConnectOne, mulai dari cara mengakses aplikasi, menggunakan '
        'fitur-fitur yang tersedia, hingga penanganan kendala umum. Setiap fitur '
        'dijelaskan dalam bentuk langkah demi langkah yang dilengkapi tangkapan '
        'layar sehingga dapat langsung diikuti oleh pengguna baru.')
    D.p('Seluruh tangkapan layar pada dokumen ini diambil dari aplikasi ConnectOne '
        'yang berjalan, menggunakan data contoh (data uji). Nama proyek, nama '
        'pengguna, dan angka yang tampak pada gambar akan berbeda dengan kondisi '
        'nyata di lingkungan produksi.')

    D.h2(2, 'Deskripsi Umum Aplikasi')
    D.p('ConnectOne adalah platform manajemen kerja internal Balai Layanan '
        'Penghubung Identitas Digital (BLPID) yang mendigitalisasi alur kerja '
        'internal organisasi dalam satu aplikasi terpadu. Ruang lingkup fungsi '
        'aplikasi meliputi:')
    D.bullets([
        '**Ruang Kerja (Workspace)**: dasbor per peran, manajemen proyek dengan '
        'kerangka kerja Scrum (epik, backlog, sprint, papan Kanban), manajemen '
        'tugas, pemantauan beban kerja, kalender agenda, laporan, dan notifikasi.',
        '**Arsip (Records)**: penyimpanan berkas (storage) dengan folder dan '
        'kendali visibilitas, registrasi dokumen resmi beserta masa berlaku, dan '
        'inventaris aset fisik.',
        '**Tata Kelola (Governance)**: manajemen perubahan (change request) '
        'berjenjang, serta tanda tangan elektronik (TTE/e-Sign) yang terintegrasi '
        'dengan layanan Balai Sertifikasi Elektronik (BSrE).',
        '**Ringkasan (Briefing)**: ringkasan harian yang menampilkan hal-hal yang '
        'memerlukan perhatian pengguna pada hari berjalan.',
        '**Administrasi**: manajemen pengguna, peran, privilege, grup pengguna, '
        'konfigurasi bot Telegram, konfigurasi e-Sign, dan log aktivitas.',
    ])

    D.h2(3, 'Pengguna Sasaran')
    D.p('Aplikasi ConnectOne digunakan oleh beberapa peran dengan cakupan akses '
        'yang berbeda. Peran menentukan menu yang tampil pada bilah navigasi dan '
        'tindakan yang boleh dilakukan.')
    D.table(['Peran', 'Deskripsi Akses'], [
        ['Administrator', 'Akses penuh terhadap konfigurasi sistem: manajemen '
         'pengguna, peran, privilege, grup pengguna, konfigurasi e-Sign dan bot '
         'Telegram, serta log aktivitas.'],
        ['Kepala Balai\n(Product Owner)', 'Memantau seluruh proyek dan tim, '
         'melihat distribusi status proyek, laporan, beban kerja, dan menyetujui '
         'permohonan perubahan.'],
        ['Kepala Seksi\n(Product Manager)', 'Mengawasi distribusi beban kerja tim, '
         'melihat seluruh proyek, laporan, dan menyetujui permohonan perubahan.'],
        ['Project Manager', 'Mengelola proyek yang menjadi tanggung jawabnya: '
         'anggota, epik, backlog, sprint, papan Kanban, dan penugasan.'],
        ['Scrum Master', 'Mengelola sprint dan memantau burndown serta velositas '
         'tim pada proyek yang diikuti.'],
        ['Staf', 'Mengerjakan tugas yang ditetapkan, mencatat waktu kerja, '
         'berkomentar, mengelola berkas dan agenda pribadi.'],
    ], widths=[1.6, 4.4])

    D.h2(4, 'Definisi dan Istilah')
    D.table(['Istilah', 'Definisi'], [
        ['TTE', 'Tanda Tangan Elektronik.'],
        ['BSrE', 'Balai Sertifikasi Elektronik, penerbit sertifikat elektronik '
         'yang digunakan pada fitur tanda tangan elektronik.'],
        ['Passphrase', 'Kata sandi sertifikat elektronik milik penanda tangan. '
         'Tidak pernah disimpan oleh aplikasi.'],
        ['NIK', 'Nomor Induk Kependudukan (16 digit), digunakan sebagai identitas '
         'penanda tangan pada layanan BSrE.'],
        ['Epik (Epic)', 'Kelompok pekerjaan besar yang menaungi beberapa item '
         'backlog.'],
        ['Backlog', 'Daftar item pekerjaan (user story, bug, fitur, tugas) yang '
         'belum dikerjakan.'],
        ['Sprint', 'Periode kerja terikat waktu dalam kerangka kerja Scrum.'],
        ['Story Point', 'Satuan estimasi ukuran pekerjaan pada sebuah item '
         'backlog.'],
        ['Burndown', 'Grafik sisa pekerjaan terhadap waktu dalam satu sprint.'],
        ['Velositas (Velocity)', 'Jumlah story point yang diselesaikan tim pada '
         'tiap sprint.'],
        ['Papan Kanban', 'Papan visual berkolom status untuk memantau dan '
         'memindahkan tugas.'],
        ['Change Request (CR)', 'Permohonan perubahan atas suatu sistem/layanan '
         'yang harus melalui peninjauan dan persetujuan.'],
        ['Penilai (Reviewer)', 'Pengguna yang meninjau dan menyetujui/menolak CR '
         'sesuai urutan yang ditentukan.'],
        ['Penandatangan', 'Pengguna yang menandatangani dokumen secara elektronik.'],
        ['Visibilitas', 'Pengaturan siapa yang dapat melihat berkas/folder pada '
         'menu Penyimpanan: Privat atau Internal.'],
        ['Privilege', 'Izin tambahan atau pencabutan izin di luar bawaan peran '
         'pengguna.'],
    ], widths=[1.6, 4.4])
    D.pagebreak()


def bab2_persyaratan(D):
    D.h1('Persyaratan Sistem')

    D.h2(1, 'Perangkat dan Peramban')
    D.bullets([
        'Komputer/laptop dengan koneksi ke jaringan internal BSSN.',
        'Peramban modern: Google Chrome, Microsoft Edge, atau Mozilla Firefox '
        'versi terbaru.',
        'Resolusi layar yang disarankan minimal 1366 × 768. Seluruh tangkapan '
        'layar pada dokumen ini diambil pada resolusi 1440 × 900.',
        'JavaScript dan cookie diaktifkan pada peramban.',
        'Aplikasi pembaca PDF untuk membuka dokumen dan laporan hasil unduhan.',
    ])

    D.h2(2, 'Akun Pengguna')
    D.bullets([
        'Akun pengguna terdaftar dengan alamat email @bssn.go.id. Akun dibuat oleh '
        'Administrator melalui menu **Kelola Pengguna**.',
        'Untuk menggunakan fitur tanda tangan elektronik: NIK 16 digit yang sudah '
        'diisi pada menu **Pengaturan**, sertifikat elektronik BSrE yang aktif, '
        'serta passphrase sertifikat tersebut.',
        'Untuk menerima notifikasi melalui Telegram: akun Telegram yang sudah '
        'ditautkan (Chat ID diisi pada menu **Pengaturan**).',
    ])

    D.h2(3, 'Batasan Berkas')
    D.table(['Fitur', 'Batasan'], [
        ['Tanda tangan elektronik (e-Sign)', 'Berkas PDF, maksimum 20 MB per '
         'dokumen.'],
        ['Distribusi e-Sign', 'Berkas PDF, maksimum 20 MB per dokumen.'],
        ['Dokumen Resmi', 'Berkas dokumen (PDF disarankan) beserta metadata masa '
         'berlaku.'],
        ['Penyimpanan (Storage)', 'Mengikuti kuota penyimpanan yang ditampilkan '
         'pada halaman Penyimpanan.'],
        ['Spesimen tanda tangan', 'Berkas gambar PNG/JPG, latar belakang '
         'transparan disarankan.'],
    ], widths=[2.4, 3.6])
    D.pagebreak()


def bab3_memulai(D):
    D.h1('Memulai')

    D.h2(1, 'Mengakses Aplikasi')
    D.p('Aplikasi ConnectOne diakses melalui peramban pada alamat berikut:')
    p = D.d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('https://connect-one.bssn.go.id')
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = NAVY
    D.p('Apabila Anda belum masuk, aplikasi akan otomatis mengarahkan ke halaman '
        'masuk (login).')

    D.h2(2, 'Masuk ke Aplikasi (Login)')
    D.steps([
        'Buka alamat aplikasi pada peramban.',
        'Pada kolom **EMAIL**, masukkan alamat email dinas Anda.',
        'Pada kolom **PASSWORD**, masukkan kata sandi. Klik ikon mata di ujung '
        'kanan kolom bila ingin menampilkan kata sandi yang diketik.',
        'Klik tombol **Sign in**. Aplikasi akan menampilkan pesan sambutan lalu '
        'mengarahkan Anda ke dasbor sesuai peran.',
    ])
    D.fig_('panduan-login.png', 'Halaman masuk (login) ConnectOne.')
    D.note('Bila email atau kata sandi tidak sesuai, aplikasi menampilkan pesan '
           '"Email atau kata sandi salah." dan Anda tetap berada di halaman masuk. '
           'Demi keamanan, pesan tidak menyebutkan bagian mana yang salah.')
    D.fig_('3.1-1_login-valid-redirect.png',
           'Hasil masuk yang berhasil, pengguna diarahkan ke dasbor sesuai perannya.')

    D.h2(3, 'Mengenal Antarmuka Aplikasi')
    D.p('Setelah masuk, seluruh halaman aplikasi memakai tata letak yang sama.')
    D.bullets([
        '**Bilah samping (sidebar) kiri**: daftar menu yang dikelompokkan menjadi '
        'RUANG KERJA (Workspace), ARSIP (Records), TATA KELOLA (Governance), dan '
        'RINGKASAN (Briefing). Menu yang tampil menyesuaikan peran Anda. Pengguna '
        'dengan hak administrasi mendapat kelompok tambahan **Admin**.',
        '**Bagian bawah sidebar**: kartu profil (nama dan peran Anda), tautan '
        '**Settings** (Pengaturan), dan tombol **Sign Out** (Keluar).',
        '**Bilah atas (header)**: tanggal hari ini, kolom **Search** (pencarian), '
        'ikon lonceng notifikasi beserta penanda jumlah yang belum dibaca, dan '
        'inisial nama Anda.',
        '**Area konten**: judul halaman, tombol tindakan utama di kanan atas, '
        'kartu statistik ringkas, lalu tabel atau daftar data.',
        'Angka pada menu **Notifications** di sidebar menunjukkan jumlah '
        'notifikasi yang belum dibaca.',
    ])
    D.fig_('3.2-3_dashboard-pm.png',
           'Tata letak umum aplikasi: sidebar menu (kiri), header (atas), dan area konten.')

    D.h2(4, 'Mengganti Bahasa Tampilan')
    D.p('Antarmuka ConnectOne tersedia dalam bahasa Inggris (bawaan) dan bahasa '
        'Indonesia. Tersedia dua cara untuk menggantinya:')
    D.bullets([
        '**Melalui header**: klik tombol **EN** atau **ID** yang berada di bilah '
        'atas, di sebelah kiri kolom pencarian. Ini cara tercepat.',
        '**Melalui Pengaturan**: klik menu **Settings** di bagian bawah sidebar, '
        'lalu pada bagian **Language** pilih **English** atau **Bahasa '
        'Indonesia**.',
    ])
    D.p('Tampilan langsung berubah tanpa perlu memuat ulang halaman. Pilihan '
        'tersimpan pada peramban Anda sehingga tetap berlaku pada kunjungan '
        'berikutnya.')
    D.note('Tangkapan layar pada dokumen ini menggunakan bahasa Inggris (bawaan '
           'aplikasi). Pada setiap langkah, nama tombol ditulis apa adanya sesuai '
           'yang tampil pada layar, diikuti padanan bahasa Indonesianya dalam '
           'tanda kurung.')

    D.h2(5, 'Keluar dari Aplikasi (Logout)')
    D.steps([
        'Klik tombol **Sign Out** (Keluar) pada bagian paling bawah sidebar.',
        'Anda akan diarahkan kembali ke halaman masuk dan sesi Anda diakhiri.',
        'Setelah keluar, membuka kembali alamat halaman dalam aplikasi (misalnya '
        'melalui tombol Back peramban) akan otomatis dialihkan ke halaman masuk.',
    ])
    D.fig_('3.1-3a_sebelum-logout.png', 'Tombol Sign Out pada bagian bawah sidebar.')
    D.fig_('panduan-setelah-logout.png',
           'Setelah keluar, akses ke halaman dalam aplikasi otomatis dialihkan ke halaman masuk.')
    D.note('Selalu klik **Sign Out** apabila Anda menggunakan komputer bersama. '
           'Menutup jendela peramban saja tidak mengakhiri sesi.')

    D.h2(6, 'Mengubah Kata Sandi')
    D.steps([
        'Klik **Settings** (Pengaturan) pada sidebar.',
        'Buka tab **Security** (Keamanan).',
        'Isi **Current Password** (Kata Sandi Saat Ini), **New Password** (Kata '
        'Sandi Baru), dan **Confirm New Password** (Konfirmasi Kata Sandi Baru).',
        'Perhatikan indikator syarat kata sandi: minimal 12 karakter, mengandung '
        'huruf besar dan kecil, minimal 1 angka, dan minimal 1 simbol. Indikator '
        '"Kata sandi cocok" muncul bila konfirmasi sudah sesuai.',
        'Klik **Change Password** (Ubah Kata Sandi). Anda akan otomatis '
        'dikeluarkan dan harus masuk kembali menggunakan kata sandi baru.',
    ])
    D.fig_('panduan-settings-security.png',
           'Halaman Pengaturan, tab Security untuk mengubah kata sandi.')
    D.note('Apabila Anda lupa kata sandi, hubungi Administrator ConnectOne untuk '
           'melakukan pengaturan ulang melalui menu Kelola Pengguna. Alamat '
           'kontak tercantum pada bab terakhir dokumen ini.')
    D.pagebreak()


def bab4_dasbor(D):
    D.h1('Dasbor dan Ringkasan Harian')
    D.p('Dasbor adalah halaman pertama yang tampil setelah Anda masuk. Isinya '
        'menyesuaikan peran, sehingga setiap peran melihat ringkasan yang relevan '
        'dengan tanggung jawabnya. Untuk kembali ke dasbor dari halaman mana pun, '
        'klik menu **Dashboard** pada sidebar.')

    D.h2(1, 'Dasbor Kepala Balai (Product Owner)')
    D.p('Menampilkan pantauan seluruh proyek dan tim.')
    D.bullets([
        'Kartu statistik: **Total Projects** (Total Proyek), **Team Members** '
        '(Anggota Tim), dan **Completed** (Selesai).',
        '**Project Status Distribution** (Distribusi Status Proyek): komposisi '
        'proyek berdasarkan status: aktif, selesai, ditunda, dan lainnya.',
        '**All Projects** (Semua Proyek): daftar seluruh proyek beserta '
        'statusnya. Klik salah satu baris untuk membuka detail proyek.',
        '**Team Members** (Anggota Tim): daftar anggota tim; klik **Manage** '
        '(Kelola) untuk menuju manajemen pengguna.',
        'Tombol **New Project** (Proyek baru) di kanan atas untuk membuat proyek '
        'baru secara langsung.',
    ])
    D.fig_('3.2-1_dashboard-po.png', 'Dasbor peran Kepala Balai (Product Owner).')

    D.h2(2, 'Dasbor Kepala Seksi (Product Manager)')
    D.p('Berfokus pada pengawasan distribusi beban kerja tim.')
    D.bullets([
        'Kartu statistik: **Total Projects** (semua proyek), **Active Projects** '
        '(proyek yang sedang berjalan), **Workload** (distribusi tim), dan '
        '**Reports** (analitik proyek).',
        'Kartu **Workload** dan **Reports** berfungsi sebagai pintasan menuju '
        'halaman Beban Kerja dan Laporan.',
        'Daftar **Projects** menampilkan rekap jumlah proyek aktif dan selesai.',
    ])
    D.fig_('3.2-2_dashboard-pdm.png', 'Dasbor peran Kepala Seksi (Product Manager).')

    D.h2(3, 'Dasbor Project Manager')
    D.p('Berfokus pada eksekusi proyek dan tugas pribadi.')
    D.bullets([
        'Kartu statistik: **Total Projects**, **My Tasks** (Tugas Saya), **In '
        'Progress** (Sedang Berjalan), dan **To Do** (Belum Dimulai).',
        'Panel **Change Request** menampilkan rekap CR per status: Draft, '
        'Submitted, Approved, Rejected, dan Implemented. Klik **View all** untuk '
        'membuka halaman Manajemen Perubahan.',
        'Panel **Projects** dan **My Tasks** menampilkan daftar ringkas; klik '
        '**View all** untuk melihat seluruh datanya.',
        'Tombol **New Sprint** (Sprint Baru) di kanan atas untuk membuat sprint.',
    ])
    D.fig_('3.2-3_dashboard-pm.png', 'Dasbor peran Project Manager.')

    D.h2(4, 'Dasbor Scrum Master')
    D.p('Berfokus pada pengelolaan sprint dan pemantauan burndown.')
    D.bullets([
        'Kartu statistik: **Total Projects**, **Active Projects**, **Completed**, '
        'dan **Workload** (pintasan untuk melihat burndown).',
        'Daftar **Active Projects** (Proyek Aktif) dengan tautan **Board & '
        'Burndown** (Papan & Burndown) pada tiap proyek.',
    ])
    D.fig_('3.2-4_dashboard-sm.png', 'Dasbor peran Scrum Master.')

    D.h2(5, 'Dasbor Staf')
    D.p('Menampilkan tugas yang ditetapkan untuk Anda dalam bentuk papan ringkas.')
    D.bullets([
        'Kartu statistik: **To Do** (Belum Dimulai), **In Progress** (Sedang '
        'Berjalan), **Completed** (Selesai), dan **Overdue** (Terlambat).',
        'Peringatan tugas terlambat beserta tombol **View** (Lihat) untuk langsung '
        'membuka tugas yang melewati tenggat.',
        'Tiga kolom tugas, **To Do**, **In Progress**, dan **Done**, berisi '
        'kartu tugas Anda. Klik kartu untuk membuka detail tugas.',
        'Tombol **All Tasks** (Semua Tugas) untuk membuka daftar tugas lengkap.',
    ])
    D.fig_('3.2-5_dashboard-staff.png', 'Dasbor peran Staf.')

    D.h2(6, 'Ringkasan Harian (Daily Brief)')
    D.p('Menu **Daily Brief** (Ringkasan Harian) pada kelompok RINGKASAN '
        'mengumpulkan seluruh hal yang memerlukan perhatian Anda pada hari '
        'berjalan, lintas modul, dalam satu halaman.')
    D.steps([
        'Klik menu **Daily Brief** pada sidebar.',
        'Bagian **Needs you today** (Memerlukan perhatian Anda hari ini) '
        'menampilkan daftar tindakan: tugas yang terlambat, tugas yang jatuh tempo '
        'hari ini, permintaan perubahan yang menunggu persetujuan Anda, serta '
        'dokumen yang menunggu tanda tangan Anda.',
        'Klik tombol pada tiap baris untuk langsung menuju objeknya, **Open** '
        '(Buka), **Review** (Tinjau), atau **Sign** (Tanda Tangani).',
        'Panel **Your agenda** (Agenda Anda) menampilkan acara kalender Anda hari '
        'ini; klik **View full calendar** (Lihat kalender lengkap) untuk membuka '
        'kalender.',
        'Panel **Waiting on others** (Menunggu orang lain) menampilkan pekerjaan '
        'yang Anda ajukan dan sedang menunggu tindakan pihak lain, lengkap dengan '
        'posisi langkahnya (misalnya "Langkah 2/3").',
    ])
    D.fig_('panduan-daily-brief.png', 'Halaman Ringkasan Harian (Daily Brief).')
    D.pagebreak()


def bab5_kalender(D):
    D.h1('Kalender dan Agenda')
    D.p('Menu **Calendar** (Kalender) digunakan untuk mencatat dan membagikan '
        'agenda kegiatan: rapat internal, kegiatan eksternal, cuti, dan kegiatan '
        'lainnya. Peserta yang diundang akan menerima notifikasi, dan sistem '
        'mengirimkan pengingat menjelang acara.')

    D.h2(1, 'Melihat Kalender')
    D.steps([
        'Klik menu **Calendar** pada sidebar.',
        'Gunakan pengalih tampilan di kanan atas untuk memilih **Day** (Hari), '
        '**Week** (Minggu), **Month** (Bulan), atau **Agenda**.',
        'Gunakan tombol panah untuk berpindah periode, dan tombol **Today** untuk '
        'kembali ke tanggal hari ini.',
        'Subjudul halaman menampilkan jumlah acara pada periode tersebut dan '
        'berapa di antaranya yang melibatkan Anda hari ini. Penanda **konflik** '
        'muncul bila terdapat acara yang waktunya bertumpang tindih.',
        'Pada tampilan bulanan, sel tanggal yang memuat banyak acara menampilkan '
        'keterangan **+n lainnya**; klik untuk melihat seluruhnya.',
        'Klik tombol **Export PDF** (Ekspor PDF) untuk mengunduh agenda dalam '
        'bentuk berkas PDF. Isi berkas mengikuti periode yang sedang '
        'ditampilkan dan hanya memuat acara yang boleh Anda lihat.',
    ])
    D.fig_('3.3.1-1a_kalender-bulanan.png', 'Kalender tampilan bulanan (Month).')
    D.fig_('3.3.1-1b_kalender-mingguan.png', 'Kalender tampilan mingguan (Week).')

    D.h2(2, 'Menambah Acara Baru')
    D.steps([
        'Pada halaman **Calendar**, klik tombol tambah acara (**New event** / '
        'Acara baru). Anda juga dapat mengeklik langsung pada sel tanggal yang '
        'dikehendaki.',
        'Isi **TITLE** (Judul) acara, kolom ini wajib diisi.',
        'Pilih **TYPE** (Jenis): **Internal**, **External** (Eksternal), **Leave** '
        '(Cuti), atau **Other** (Lainnya).',
        'Pilih **VISIBILITY** (Visibilitas): **Public** (Publik) agar acara '
        'terlihat oleh pengguna lain, atau **Private** (Privat) agar hanya '
        'terlihat oleh Anda dan peserta yang diundang.',
        'Tentukan **START** (Mulai) dan **END** (Selesai). Centang **All day** '
        '(Sepanjang hari) bila acara berlangsung seharian penuh; bila tidak, isi '
        '**START TIME** dan **END TIME**.',
        'Isi **LOCATION** (Lokasi), misalnya "Daring" atau "Ruang Rapat A".',
        'Isi **DESCRIPTION** (Deskripsi) sebagai catatan tambahan bila diperlukan.',
        'Pilih peserta pada bagian peserta acara (lihat sub-bab 5.5).',
        'Klik **Save event** (Simpan acara). Aplikasi menampilkan pesan '
        '"Acara berhasil ditambahkan!" dan acara langsung muncul pada kalender.',
    ])
    D.fig_('3.3.2-1a_form-tambah-kegiatan.png', 'Formulir penambahan acara baru.')
    D.fig_('3.3.2-1b_kegiatan-muncul-di-kalender.png',
           'Acara yang baru dibuat muncul pada tanggal yang dipilih.')
    D.note('Judul dan tanggal mulai wajib diisi. Bila salah satunya kosong, '
           'aplikasi menampilkan pesan "Judul dan tanggal mulai wajib diisi" dan '
           'acara tidak tersimpan.')

    D.h2(3, 'Melihat Detail dan Mengubah Acara')
    D.steps([
        'Klik acara pada kalender untuk membuka panel **Event details** (Detail '
        'acara). Panel menampilkan TANGGAL, WAKTU, LOKASI, PESERTA, dan STATUS '
        '(Akan Datang / Berlangsung / Selesai).',
        'Klik ikon ubah (**Edit**) pada panel detail.',
        'Ubah kolom yang diperlukan pada formulir yang tampil.',
        'Klik **Save event** (Simpan acara). Aplikasi menampilkan pesan '
        '"Acara berhasil diperbarui!" dan perubahan langsung terlihat di kalender.',
    ])
    D.fig_('3.3.2-2a_detail-kegiatan.png', 'Panel detail acara.')
    D.fig_('3.3.2-2b_form-edit-kegiatan.png', 'Formulir pengubahan acara.')
    D.fig_('3.3.2-2c_perubahan-tersimpan.png', 'Perubahan tersimpan dan tampil pada kalender.')

    D.h2(4, 'Menghapus Acara')
    D.steps([
        'Klik acara pada kalender untuk membuka panel detail.',
        'Klik ikon hapus (**Delete**).',
        'Konfirmasikan penghapusan pada dialog yang muncul.',
        'Aplikasi menampilkan pesan "Acara berhasil dihapus!" dan acara hilang '
        'dari kalender.',
    ])
    D.fig_('3.3.2-3a_detail-sebelum-hapus.png', 'Detail acara sebelum dihapus.')
    D.fig_('3.3.2-3b_setelah-hapus.png', 'Kalender setelah acara dihapus.')
    D.note('Penghapusan acara bersifat permanen dan tidak dapat dibatalkan. '
           'Peserta yang sudah diundang tidak lagi melihat acara tersebut.',
           label='Perhatian')

    D.h2(5, 'Mengundang Peserta')
    D.p('Peserta dapat dipilih satu per satu (perorangan) atau sekaligus melalui '
        'grup pengguna yang sudah dibuat Administrator.')
    D.steps([
        'Pada formulir acara, buka bagian peserta.',
        'Gunakan kolom **Search by name** (Cari berdasarkan nama) untuk menyaring '
        'daftar pengguna.',
        'Centang nama pengguna yang akan diundang. Jumlah yang dipilih tampil '
        'sebagai keterangan "**n dipilih**".',
        'Untuk mengundang satu unit sekaligus, pilih grup pada bagian **Participant '
        'Groups** (Grup Peserta). Keterangan jumlah anggota grup tampil di '
        'sampingnya.',
        'Simpan acara. Seluruh peserta yang dipilih tercatat pada panel detail '
        'acara dan menerima notifikasi undangan.',
    ])
    D.fig_('3.3.3-1a_form-pilih-banyak-peserta.png',
           'Memilih beberapa peserta sekaligus pada formulir acara.')
    D.fig_('3.3.3-1c_daftar-peserta-terdaftar.png',
           'Daftar peserta yang terdaftar pada detail acara.')

    D.h2(6, 'Acara Publik dan Acara Privat')
    D.p('Pengaturan **VISIBILITY** (Visibilitas) menentukan siapa yang dapat '
        'melihat sebuah acara.')
    D.table(['Visibilitas', 'Siapa yang dapat melihat'], [
        ['Public (Publik)', 'Seluruh pengguna aplikasi dapat melihat acara pada '
         'kalender masing-masing.'],
        ['Private (Privat)', 'Hanya pembuat acara dan peserta yang diundang. '
         'Pengguna lain tidak melihat acara ini sama sekali.'],
    ], widths=[1.8, 4.2])
    D.fig_('3.3.3-2a_form-acara-privat.png', 'Menetapkan acara sebagai Private (Privat).')
    D.fig_('3.3.3-2b_pemilik-melihat-kedua-acara.png',
           'Kalender pemilik acara, acara publik dan privat keduanya terlihat.')
    D.fig_('3.3.3-2c_pengguna-lain-bukan-peserta.png',
           'Kalender pengguna lain yang bukan peserta, acara privat tidak ditampilkan.')

    D.h2(7, 'Mencatat Notulensi dan Tindak Lanjut')
    D.p('Setelah acara berlangsung, hasil pembahasannya dapat dicatat langsung '
        'pada acara tersebut.')
    D.steps([
        'Buka panel detail acara.',
        'Klik **+ Tambah laporan** (Add report).',
        'Isi bagian **NOTULENSI** (Minutes), **HASIL PEMBAHASAN** (Outcomes), dan '
        '**TINDAK LANJUT** (Follow up).',
        'Simpan. Catatan akan tampil pada panel detail acara dan dapat dibaca oleh '
        'peserta acara.',
    ])

    D.h2(8, 'Notifikasi dan Pengingat Agenda')
    D.bullets([
        'Saat acara baru dibuat dan Anda ditambahkan sebagai peserta, Anda '
        'menerima notifikasi di dalam aplikasi serta melalui Telegram (bila Chat '
        'ID sudah ditautkan).',
        'Sistem mengirimkan pengingat otomatis **satu hari sebelum** acara (H-1) '
        'dan **pada hari pelaksanaan** acara (H-0).',
        'Jenis notifikasi yang diterima dapat diatur pada menu **Settings → '
        'Notifications** (lihat Bab 14).',
    ])
    D.fig_('3.3.4-1_bukti-notifikasi-telegram.png',
           'Notifikasi acara kalender yang diterima melalui Telegram.')
    D.fig_('3.3.4-2_bukti-pengingat-h1-h0.png',
           'Pengingat otomatis H-1 dan H-0 untuk agenda kegiatan.')
    D.pagebreak()
