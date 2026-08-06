#!/usr/bin/env python3
"""Pass kedua: penanda status dokumen, isi metadata, dan perbaiki kesimpulan."""
import copy, json, os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE = os.path.dirname(os.path.abspath(__file__))
DOC = '/home/ymjsty/dev/agrawork/UAT_ConnectOne_v3_HASIL-VERIFIKASI-TEKNIS.docx'
results = {r['id']: r for r in json.load(open(os.path.join(BASE, 'ALL-RESULTS.json')))}

RED = RGBColor(0xC0, 0x1B, 0x1B)
NAVY = RGBColor(0x14, 0x40, 0x6A)


def insert_before(paragraph, text, bold=False, size=10, color=None, italic=False):
    new_p = copy.deepcopy(paragraph._p)
    paragraph._p.addprevious(new_p)
    from docx.text.paragraph import Paragraph
    p = Paragraph(new_p, paragraph._parent)
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    run = p.add_run(text)
    run.font.bold = bold
    run.font.size = Pt(size)
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p


def main():
    doc = Document(DOC)
    paras = doc.paragraphs

    # ---------- 1. Penanda status di awal dokumen ----------
    first = next((p for p in paras if p.text.strip().startswith('BAB I')), paras[0])
    fails = sorted([r for r in results.values() if not r['verdict'].startswith('PASS')], key=lambda r: r['id'])
    total = len(results)
    npass = sum(1 for r in results.values() if r['verdict'].startswith('PASS'))

    insert_before(first, '')
    insert_before(first,
        'Dokumen ini BELUM merupakan hasil User Acceptance Test resmi.',
        bold=True, size=12, color=RED)
    insert_before(first,
        'Isian hasil dan bukti tangkapan layar pada dokumen ini berasal dari VERIFIKASI TEKNIS INTERNAL (pre-UAT) '
        'yang dijalankan secara otomatis oleh tim pengembang pada lingkungan pengujian lokal, bukan dari pengujian '
        'manual oleh perwakilan pengguna di lingkungan BLPID sebagaimana dinyatakan pada Bab II. '
        'Dokumen ini disusun sebagai DRAFT dan bahan persiapan sebelum UAT resmi dilaksanakan.',
        size=10)
    insert_before(first, '')
    if fails:
        insert_before(first,
            f'Ringkasan hasil verifikasi teknis: {total} skenario diuji, {npass} sesuai ekspektasi, '
            f'{len(fails)} TIDAK sesuai ekspektasi dan perlu perbaikan sebelum go-live '
            f'({", ".join(r["id"] for r in fails)}). Rincian pada Bab V.',
            size=10, bold=True)
    else:
        insert_before(first,
            f'Ringkasan hasil verifikasi teknis: {total} skenario diuji dan seluruhnya sesuai ekspektasi. '
            'Pada putaran verifikasi pertama ditemukan 5 ketidaksesuaian (3.4.4-2, 3.6.2-1, 3.8.1-4, 4.1-3, dan '
            '4.2-5); seluruhnya telah diperbaiki dan diuji ulang, dan catatan hasil pada bab berikut memuat '
            'penjelasan penyebab beserta bukti pengujian ulangnya. Butir tindak lanjut yang masih terbuka '
            'tercantum pada Bab V.',
            size=10, bold=True)
    insert_before(first, '')
    insert_before(first,
        'Bab VI (Pernyataan Penerimaan) sengaja DIBIARKAN KOSONG. Keputusan diterima atau ditolak beserta '
        'penandatanganan sepenuhnya merupakan kewenangan perwakilan BLPID setelah UAT resmi dilaksanakan.',
        size=10, italic=True)
    insert_before(first, '')

    # ---------- 2. Metadata pelaksanaan (tabel 0) ----------
    t0 = doc.tables[0]
    for row in t0.rows:
        key = row.cells[0].text.strip().lower()
        if key.startswith('periode'):
            row.cells[1].text = '30 Juli – 1 Agustus 2026 (verifikasi teknis internal / pre-UAT, termasuk pengujian ulang setelah perbaikan)'
        elif key.startswith('lokasi'):
            row.cells[1].text = 'Lingkungan pengujian lokal tim pengembang (bukan lingkungan BLPID)'
        elif key.startswith('url'):
            row.cells[1].text = 'http://localhost:3000 (frontend) · http://localhost:8000 (API gateway)'
        for c in row.cells:
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)

    # ---------- 3. Perbaiki kalimat kesimpulan 5.2 ----------
    for p in doc.paragraphs:
        if 'seluruh 74 skenario' in p.text or '74 skenario uji' in p.text:
            for r in list(p.runs):
                r._element.getparent().remove(r._element)
            if fails:
                teks = (
                    f'Berdasarkan verifikasi teknis internal terhadap {total} skenario uji dari 11 modul, sebanyak '
                    f'{npass} skenario berperilaku sesuai ekspektasi dan {len(fails)} skenario TIDAK memenuhi ekspektasi. '
                    'Dengan demikian sistem ConnectOne BELUM sepenuhnya memenuhi kebutuhan fungsional yang ditetapkan '
                    'dalam dokumen pengadaan dan KAK. Temuan berikut perlu diperbaiki terlebih dahulu: '
                    + '; '.join(f'{r["id"]} {r["scenario"]}' for r in fails) + '. '
                    'Kesimpulan akhir mengenai kelayakan operasional tetap menjadi kewenangan perwakilan BLPID '
                    'melalui pelaksanaan UAT resmi.')
            else:
                teks = (
                    f'Berdasarkan verifikasi teknis internal terhadap {total} skenario uji dari 11 modul, seluruh '
                    f'{npass} skenario berperilaku sesuai ekspektasi. Putaran verifikasi pertama menemukan 5 '
                    'ketidaksesuaian pada modul beban kerja, tanda tangan elektronik multi pihak, manajemen pengguna, '
                    'sprint, dan manajemen dokumen; seluruhnya telah diperbaiki dan diuji ulang hingga sesuai '
                    'ekspektasi. Dengan demikian sistem ConnectOne memenuhi kebutuhan fungsional yang diuji pada '
                    'lingkungan pengujian internal, dengan sejumlah butir tindak lanjut yang tercantum pada '
                    'sub-bab 5.3. Kesimpulan akhir mengenai kelayakan operasional tetap menjadi kewenangan '
                    'perwakilan BLPID melalui pelaksanaan UAT resmi.')
            run = p.add_run(teks)
            run.font.size = Pt(11)
            break

    # ---------- 4. Catatan metode pada Bab II ----------
    for p in doc.paragraphs:
        if p.text.strip().startswith('Pengujian dilakukan secara manual oleh tim penguji'):
            note = insert_before(p, '')
            insert_before(p,
                'Catatan pelaksanaan: metode di bawah ini menggambarkan rencana UAT resmi. Pengisian hasil pada '
                'dokumen versi ini dilakukan melalui pengujian otomatis berbasis peramban (Playwright) oleh tim '
                'pengembang, dengan verifikasi silang ke basis data dan log layanan untuk setiap skenario.',
                size=9.5, italic=True, color=NAVY)
            break

    doc.save(DOC)
    print('  selesai:', DOC)
    print(f'  {total} skenario | {npass} sesuai | {len(fails)} tidak sesuai')


if __name__ == '__main__':
    main()
