#!/usr/bin/env python3
"""Susun Buku Panduan Penggunaan ConnectOne lengkap dengan langkah + screenshot.

Screenshot diambil dari uat-evidence/shots/ (bukti verifikasi teknis pra-UAT).
Jalankan: python3 build_panduan.py
"""
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = os.path.dirname(os.path.abspath(__file__))
SHOTS = os.path.join(BASE, 'shots')
OUT = '/home/ymjsty/dev/agrawork/Buku_Panduan_Penggunaan_ConnectOne.docx'

NAVY = RGBColor(0x1B, 0x2C, 0x4B)
GREY = RGBColor(0x55, 0x55, 0x55)
DIM = RGBColor(0x77, 0x77, 0x77)
IMG_W = Inches(5.9)

# ---------------------------------------------------------------- primitives

def shade(el, hexcolor):
    tc = OxmlElement('w:shd')
    tc.set(qn('w:val'), 'clear')
    tc.set(qn('w:color'), 'auto')
    tc.set(qn('w:fill'), hexcolor)
    el.append(tc)


def cell_shade(cell, hexcolor):
    shade(cell._tc.get_or_add_tcPr(), hexcolor)


def set_borders(table, color='BFBFBF', sz=4):
    tbl = table._tbl
    pr = tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        e = OxmlElement(f'w:{edge}')
        e.set(qn('w:val'), 'single')
        e.set(qn('w:sz'), str(sz))
        e.set(qn('w:space'), '0')
        e.set(qn('w:color'), color)
        borders.append(e)
    pr.append(borders)


def cell_text(cell, text, bold=False, size=10, color=None, align=None):
    cell.text = ''
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    if align is not None:
        p.alignment = align
    first = True
    for line in str(text).split('\n'):
        if not first:
            p = cell.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(3)
            if align is not None:
                p.alignment = align
        r = p.add_run(line)
        r.bold = bold
        r.font.size = Pt(size)
        if color is not None:
            r.font.color.rgb = color
        first = False


class Doc:
    def __init__(self):
        d = Document()
        # A4 + margin 2,5 cm
        s = d.sections[0]
        s.page_width = Emu(7560310)
        s.page_height = Emu(10692130)
        for attr in ('left_margin', 'right_margin', 'top_margin', 'bottom_margin'):
            setattr(s, attr, Inches(1))

        normal = d.styles['Normal']
        normal.font.name = 'Calibri'
        normal.font.size = Pt(11)
        normal.paragraph_format.space_after = Pt(6)
        normal.paragraph_format.line_spacing = 1.15

        for name, size, color in (
            ('Heading 1', 17, NAVY), ('Heading 2', 14, NAVY), ('Heading 3', 12, NAVY),
        ):
            st = d.styles[name]
            st.font.name = 'Calibri'
            st.font.size = Pt(size)
            st.font.color.rgb = color
            st.font.bold = True
            st.paragraph_format.space_before = Pt(14 if name == 'Heading 1' else 10)
            st.paragraph_format.space_after = Pt(6)
        cap = d.styles['Caption']
        cap.font.size = Pt(9)
        cap.font.italic = True
        cap.font.color.rgb = GREY

        self.d = d
        self.chapter = 0
        self.fig = 0
        self.missing = []
        self._footer()

    def _footer(self):
        f = self.d.sections[0].footer
        p = f.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run('Buku Panduan Penggunaan ConnectOne  ·  BLPID  ·  Halaman ')
        r.font.size = Pt(8)
        r.font.color.rgb = DIM
        # field PAGE
        fld = OxmlElement('w:fldSimple')
        fld.set(qn('w:instr'), 'PAGE')
        run = OxmlElement('w:r')
        rpr = OxmlElement('w:rPr')
        sz = OxmlElement('w:sz'); sz.set(qn('w:val'), '16'); rpr.append(sz)
        run.append(rpr)
        t = OxmlElement('w:t'); t.text = '1'; run.append(t)
        fld.append(run)
        p._p.append(fld)

    # ---- blocks -----------------------------------------------------------
    def h1(self, text):
        self.chapter += 1
        self.fig = 0
        p = self.d.add_paragraph(f'{self.chapter}. {text}', style='Heading 1')
        return p

    def h2(self, num, text):
        return self.d.add_paragraph(f'{self.chapter}.{num} {text}', style='Heading 2')

    def h3(self, text):
        return self.d.add_paragraph(text, style='Heading 3')

    def p(self, text, size=11, italic=False, color=None, space_after=6,
          align=WD_ALIGN_PARAGRAPH.JUSTIFY):
        par = self.d.add_paragraph()
        par.paragraph_format.space_after = Pt(space_after)
        if align is not None:
            par.alignment = align
        for run_text, bold in parse_bold(text):
            r = par.add_run(run_text)
            r.bold = bold
            r.italic = italic
            r.font.size = Pt(size)
            if color is not None:
                r.font.color.rgb = color
        return par

    def steps(self, items):
        for i, it in enumerate(items, 1):
            par = self.d.add_paragraph()
            pf = par.paragraph_format
            pf.left_indent = Inches(0.42)
            pf.first_line_indent = Inches(-0.42)
            pf.space_after = Pt(4)
            par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            r = par.add_run(f'{i}.\t')
            r.bold = True
            for run_text, bold in parse_bold(it):
                rr = par.add_run(run_text)
                rr.bold = bold

    def bullets(self, items):
        for it in items:
            par = self.d.add_paragraph()
            pf = par.paragraph_format
            pf.left_indent = Inches(0.42)
            pf.first_line_indent = Inches(-0.2)
            pf.space_after = Pt(3)
            par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            par.add_run('•\t')
            for run_text, bold in parse_bold(it):
                rr = par.add_run(run_text)
                rr.bold = bold

    def note(self, text, label='Catatan'):
        t = self.d.add_table(rows=1, cols=1)
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_borders(t, 'C9D4E4')
        c = t.cell(0, 0)
        cell_shade(c, 'F1F5FA')
        c.text = ''
        par = c.paragraphs[0]
        par.paragraph_format.space_before = Pt(4)
        par.paragraph_format.space_after = Pt(4)
        par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r = par.add_run(f'{label}: ')
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = NAVY
        for run_text, bold in parse_bold(text):
            rr = par.add_run(run_text)
            rr.bold = bold
            rr.font.size = Pt(10)
        self.d.add_paragraph().paragraph_format.space_after = Pt(2)

    def fig_(self, filename, caption):
        """Sisipkan screenshot + keterangan bernomor otomatis."""
        self.fig += 1
        label = f'Gambar {self.chapter}.{self.fig}'
        path = os.path.join(SHOTS, filename) if filename else None
        if path and os.path.exists(path):
            par = self.d.add_paragraph()
            par.alignment = WD_ALIGN_PARAGRAPH.CENTER
            par.paragraph_format.space_before = Pt(6)
            par.paragraph_format.space_after = Pt(2)
            par.add_run().add_picture(path, width=IMG_W)
        else:
            if filename:
                self.missing.append(filename)
            t = self.d.add_table(rows=1, cols=1)
            t.alignment = WD_TABLE_ALIGNMENT.CENTER
            set_borders(t, 'BFBFBF')
            c = t.cell(0, 0)
            cell_shade(c, 'FAFAFA')
            cell_text(c, '\n[ Tangkapan layar untuk bagian ini belum tersedia, '
                         'sisipkan di sini ]\n', size=9, color=DIM,
                      align=WD_ALIGN_PARAGRAPH.CENTER)
        cp = self.d.add_paragraph(style='Caption')
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_after = Pt(10)
        cp.add_run(f'{label}  {caption}')

    def table(self, headers, rows, widths=None, header_fill='1B2C4B'):
        t = self.d.add_table(rows=1, cols=len(headers))
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_borders(t)
        for i, h in enumerate(headers):
            c = t.rows[0].cells[i]
            cell_shade(c, header_fill)
            cell_text(c, h, bold=True, size=10, color=RGBColor(0xFF, 0xFF, 0xFF))
        for row in rows:
            cells = t.add_row().cells
            for i, v in enumerate(row):
                cell_text(cells[i], v, size=10)
        if widths:
            for r in t.rows:
                for i, w in enumerate(widths):
                    r.cells[i].width = Inches(w)
        self.d.add_paragraph().paragraph_format.space_after = Pt(2)
        return t

    def kv_table(self, rows, w1=1.8, w2=4.2):
        t = self.d.add_table(rows=0, cols=2)
        set_borders(t)
        for k, v in rows:
            cells = t.add_row().cells
            cell_shade(cells[0], 'F1F5FA')
            cell_text(cells[0], k, bold=True, size=10)
            cell_text(cells[1], v, size=10)
            cells[0].width = Inches(w1)
            cells[1].width = Inches(w2)
        self.d.add_paragraph().paragraph_format.space_after = Pt(2)
        return t

    def pagebreak(self):
        self.d.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    def save(self):
        self.d.save(OUT)


def parse_bold(text):
    """Pisah teks ber-**tebal** menjadi daftar (potongan, bold)."""
    out, buf, bold = [], '', False
    i = 0
    while i < len(text):
        if text[i:i + 2] == '**':
            if buf:
                out.append((buf, bold))
                buf = ''
            bold = not bold
            i += 2
        else:
            buf += text[i]
            i += 1
    if buf:
        out.append((buf, bold))
    return out or [('', False)]


# ------------------------------------------------------------------- main

def main():
    D = Doc()
    import panduan_isi_a as A
    import panduan_isi_b as B
    import panduan_isi_c as C

    A.cover(D)
    A.info_dokumen(D)
    A.bab1_pendahuluan(D)
    A.bab2_persyaratan(D)
    A.bab3_memulai(D)
    A.bab4_dasbor(D)
    A.bab5_kalender(D)
    B.bab6_proyek(D)
    B.bab7_tugas(D)
    B.bab8_beban_kerja(D)
    B.bab9_laporan(D)
    B.bab10_change_request(D)
    C.bab11_tte(D)
    C.bab12_arsip(D)
    C.bab13_notifikasi(D)
    C.bab14_pengaturan(D)
    C.bab15_admin(D)
    C.bab16_troubleshooting(D)
    C.bab17_faq(D)
    C.bab18_kontak(D)

    D.save()
    print(f'[ok] {OUT}')
    if D.missing:
        print('[!] berkas screenshot tidak ditemukan:')
        for m in sorted(set(D.missing)):
            print('    -', m)


if __name__ == '__main__':
    main()
